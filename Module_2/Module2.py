import os
import time 
import pyautogui
from pynput import keyboard,mouse
from docx import Document
from docx.shared import Inches
from winotify import Notification
from pathlib import Path
from docx2pdf import convert

last_activity_time=time.time()
Notification_pause=3
setup_time=8
navigation_time=10

def Show_notification(title,message):
     toast = Notification(
          app_id="Lab_Docs_Screenshot_Automation",
          title=title,
          msg=message,
          duration='short')
     toast.show()

def is_word_file_available(filepath):
    if not os.path.exists(filepath):
        return True
    try:
        with open(filepath,"a") as f:
            f.close()
        return True
    except IOError:
        return False

def regionCoordinates():
    Show_notification("Notification","Move your mouse to the top-left corner of the region and wait...")
    time.sleep(setup_time)
    top_left_x,top_left_y=pyautogui.position()
    print(f"Top-left coordinates: ({top_left_x},{top_left_y})")

    time.sleep(setup_time)
    Show_notification("Notification","Move your mouse to the bottom-right corner and wait..")
    time.sleep(setup_time)
    bottom_right_x,bottom_right_y=pyautogui.position()
    Show_notification("Status",f"Top-left coordinates: ({top_left_x},{top_left_y}),\nBottom-right co ordinates:({bottom_right_x},{bottom_right_y})")

    print(f"Bottom-right co ordinates:({bottom_right_x},{bottom_right_y})")

    width = bottom_right_x - top_left_x
    height = bottom_right_y - top_left_y
    
    Show_notification("Notification","Region Coordionates are taken")
    time.sleep(Notification_pause)
    region=(top_left_x,top_left_y,width,height)
    return region

""" these Updates the last activity time on user actions."""
def on_key_press(key):
     global last_activity_time
     last_activity_time = time.time()

def on_mouse_click(x,y,button,press):
     global last_activity_time
     if press:
        last_activity_time = time.time()
    
def  on_mouse_move(x,y):
     global last_activity_time
     last_activity_time = time.time()

def  setup_listeners():
     keyboard_listener = keyboard.Listener(on_press=on_key_press)
     mouse_listener=mouse.Listener(on_click=on_mouse_click,on_move=on_mouse_move)
     keyboard_listener.start()
     mouse_listener.start()
     return keyboard_listener,mouse_listener
     
#if run is constant you can automate the running program with this
# def runbuttonCoordinates():
#     print("Move your mouse over the 'run button'")
#     time.sleep(5)
#     click_x,click_y = pyautogui.position()
#     print(f"Run Button coordinates:({click_x},{click_y})")
#     return click_x,click_y

# rb_x, rb_y = runbuttonCoordinates()
# pyautogui.click(rb_x,rb_y)

def handle_screenshot_logic(placeholder, idle_timeOut_seconds, setup_time):
    global last_activity_time
    original_screenshot_path = "Screenshot.png" 
    
    Show_notification("Instructions", f"Ready for Screenshot for: {placeholder}. Please navigate to the correct window.")
    time.sleep(navigation_time)
    region = regionCoordinates() 
    
    screenshot_taken = False
    
    while not screenshot_taken:
        time.sleep(1) 
        

        idle_time = time.time() - last_activity_time
        

        if idle_time >= idle_timeOut_seconds:
            Show_notification("Alert!", "No activity detected. Taking Screenshot.")
            time.sleep(setup_time) 
            
           
            Screenshot = pyautogui.screenshot(region=region)
            Screenshot.save(original_screenshot_path)
            
           
            last_activity_time = time.time()
            screenshot_taken = True
        
    return original_screenshot_path
     
     


def Lab_Docs_Screenshot_Automation(file_path,placeholders,idle_timeOut_seconds,img_width):
    global last_activity_time
    Show_notification("Attention",f"Please Make sure the Placeholders in List placeholders and file '{file_path}' are same.")
    time.sleep(Notification_pause)

    if Path(file_path).exists():
            if is_word_file_available(file_path):
                    Show_notification("Status",f"The file '{file_path}' is available.You can begin your automation")
            else:
                    Show_notification("Problem Occurred",f"The file '{file_path}' is currently in use by another application./n Please close file and restart Automation")
            document = Document(file_path)
    else:
        Show_notification("Error Report",f"The file '{file_path}' is doesn't exist ./n Please close file and restart Automation")
        exit()
            
    Show_notification("Instructions",f"Automation is about to begin ,\nPlease ensure not open the file '{file_path}' in the middle of the process")

    keyboard_listener,mouse_listener = setup_listeners()
    Show_notification("Attention",f"Monitoring activity. Idle timeout set to '{idle_timeOut_seconds}' seconds")
    time.sleep(Notification_pause)

    
    try:
        for placeholder in placeholders:

            original_screenshot_path = handle_screenshot_logic(placeholder, idle_timeOut_seconds, setup_time)
            img_to_add = original_screenshot_path
            
            replaced=False

            for paragraph in document.paragraphs:
                if replaced:
                     break
                if placeholder in paragraph.text:
                    for run in paragraph.runs:
                        if placeholder in run.text: 
                                run.text=run.text.replace(placeholder,"")
                                run.add_picture(img_to_add,width=Inches(img_width))
                                print(f"Placeholder '{placeholder}' replaced with image in a run.")
                                Show_notification("Status",f"Placeholder '{placeholder}' replaced with image.")
                                time.sleep(navigation_time)
                                replaced=True
                                break
            if not replaced:
                        Show_notification("Error Report",f"Placeholder '{placeholder}' not found in any run/paragraph. Check file structure.")
                        time.sleep(navigation_time)
                        
                    
                
    except KeyboardInterrupt:
          Show_notification("Notification","ScreenShot Automation Stopped by User.")
    finally:
        document.save("Output.docx")
        # convert("Output.docx","LabManual.pdf")
        Show_notification("Success","Output document saved successfully")                
        if os.path.exists(original_screenshot_path):
             os.remove(original_screenshot_path)
        try:
            input("Press Enter to exit....") 
        except EOFError:
            pass


if __name__ == "__main__":
     file_path="labdocs.docx"
     placeholders=["Lab1_code","Output1","Lab2_code","Output2","Lab3_code","Output3"]
     idle_timeOut_seconds = 20
     img_width=6.53
     Lab_Docs_Screenshot_Automation(file_path,placeholders,idle_timeOut_seconds,img_width)

