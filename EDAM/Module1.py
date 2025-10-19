import os
import time
import sys
import pyautogui
from PIL import Image
from pathlib import Path
from winotify import Notification
from pynput import keyboard, mouse
from docx import Document
from docx.shared import Inches
from docx2pdf import convert

setup_time = 8
navigation_time = 15
Notification_pause = 3
last_activity_time = time.time() 


def Show_notifications(title,message):
    toast = Notification(
        app_id="ScreenShot Automation",
        title=title,
        msg=message,
        duration="short" )
    toast.show()
    
def is_word_file_available(filepath):
    """Checks if the word file is closed and available for writing."""
    if not os.path.exists(filepath):
        return True
    try:
        # Attempt to open in append mode to check write access
        with open(filepath,"a") as f:
            f.close()
        return True
    except IOError:
        return False
    
def rotate_image(image_path,angle):
    """Rotates an image and saves the rotated version."""
    try:
        img = Image.open(image_path)
        rotated_image=img.rotate(angle,expand = True)
        
        # Create a unique path for the rotated image
        rotated_image_path=f"rotated_{angle}_{image_path}"
        rotated_image.save(rotated_image_path)
        return rotated_image_path
    except Exception as e:
        print(f"Error rotating/saving image: {e}")
        Show_notifications("Error","Image rotation failed.")
        return None
 

def regionCoordinates():
    #print("Move your mouse to the top-left corner of the region and wait...")
    #time.sleep(5)
    Show_notifications("Region Setup","Move your mouse to the top-left corner of the region and wait...")
    time.sleep(setup_time)
    top_left_x,top_left_y=pyautogui.position()
    print(f"Top-left coordinates: ({top_left_x},{top_left_y})")

    #print("Move your mouse to the bottom-right corner and wait..")
    time.sleep(setup_time)
    Show_notifications("Region Setup","Move your mouse to the bottom-right corner and wait..")
    time.sleep(setup_time)
    bottom_right_x,bottom_right_y=pyautogui.position()
    Show_notifications("Status",f"Top-left coordinates: ({top_left_x},{top_left_y}),\nBottom-right co ordinates:({bottom_right_x},{bottom_right_y})")
    print(f"Bottom-right co ordinates:({bottom_right_x},{bottom_right_y})")

    width = bottom_right_x - top_left_x
    height = bottom_right_y - top_left_y
    
    Show_notifications("Region Status ","Region Coordionates are taken.")
    time.sleep(setup_time)
    region_Coordinates=(top_left_x,top_left_y,width,height)
    return region_Coordinates


def navigationCoordinates():  
    Show_notifications("Navigation Setup:","Move your mouse over the 'NEXT' button or slide area and wait.." )
    time.sleep(navigation_time)
    click_x,click_y = pyautogui.position()
    print(f"Navigation click coordinates:({click_x},{click_y})")
    Show_notifications("Navigation Setup:",f"Next Page click Coordinate is set  at ({click_x},{click_y})")
    return click_x,click_y


def on_key_press(key):
    """Updates the last activity time on key press."""
    global last_activity_time
    last_activity_time=time.time()
   
def on_mouse_click(x,y,button,press):
    """Updates the last activity time on mouse click."""
    global last_activity_time
    if press:
        last_activity_time = time.time()
       
def on_mouse_move(x,y):
    """Updates the last activity time on mouse movement."""
    global last_activity_time
    last_activity_time=time.time()

def setup_listeners():
    """Initializes and starts the keyboard and mouse activity listeners."""
    keyboard_listener= keyboard.Listener(on_press=on_key_press)
    mouse_listener=mouse.Listener(on_click=on_mouse_click,on_move=on_mouse_move)
    keyboard_listener.start()
    mouse_listener.start()
    return keyboard_listener,mouse_listener

    
def Screenshot_Automation(file_path,idle_timeOut_seconds,rotation_angle,docx_width,docs_height):

    global last_activity_time
    original_screenshot_path="screenshot.png"
    Rotated_image = None 
    

    if Path(file_path).exists():
        if is_word_file_available(file_path):
                Show_notifications("Status",f"The file {file_path} is available. You can begin your automation")
                time.sleep(Notification_pause)
        else:
                Show_notifications("Error Report",f"The file {file_path} is currently in use by another application.\nTerminating the Program")
                time.sleep(Notification_pause)
                Show_notifications("Follow the instruction",f"Please close {file_path} and Restart Automation.")
                time.sleep(Notification_pause)
                sys.exit()        
        document = Document(file_path)
    else:
        document=Document()


    Show_notifications("Attention:",f"Starting the Process,Please Ensure Not to Open file {file_path} in middle of Automation")
    time.sleep(Notification_pause)
    

    region = regionCoordinates()
    clickx,clicky = navigationCoordinates()
    

    keyboard_listener,mouse_listener = setup_listeners()


    Show_notifications("Attention",f"Monitoring activity. Idle timeout set to {idle_timeOut_seconds} seconds")
    time.sleep(Notification_pause)

    try:
        while True:
            idle_time = time.time() - last_activity_time
            if idle_time > idle_timeOut_seconds:
                Show_notifications("Alert!","No activity detected. Taking Screenshot.")
                time.sleep(setup_time)
    
                screenshot = pyautogui.screenshot(region=region)
                screenshot.save(original_screenshot_path)
                
                # Assume original screenshot path is used by default
                image_to_add = original_screenshot_path
                
                if rotation_angle != 0:
                    Rotated_image = rotate_image(original_screenshot_path, rotation_angle)
                    
                    # Check if rotation was successful before overwriting image_to_add
                    if Rotated_image:
                        image_to_add = Rotated_image
                        
                # Check if a valid image path exists to be added
                if image_to_add and os.path.exists(image_to_add):
                    document.add_picture(image_to_add, width=Inches(docx_width), height=Inches(docs_height))
                    Show_notifications("Status","Screenshot Added to word file.\nMoving to Next Page")
                else:
                    # This block handles failed rotation or other capture issues
                    Show_notifications("Error", "Could not capture or process image. Skipping.")

                time.sleep(Notification_pause)

                pyautogui.click(clickx,clicky)
                last_activity_time = time.time() 
                time.sleep(setup_time)
                

    except KeyboardInterrupt:
        pass
    finally:
        Show_notifications("Status","ScreenShot Automation Stopped by User.")
        time.sleep(Notification_pause)
        keyboard_listener.stop()
        mouse_listener.stop()
        document.save(file_path)
        #convert(file_path,"newdocs1.pdf")
        Show_notifications("Status","Document saved successfully.")
        
        # Only remove the rotated file if the path is defined AND the file exists
        if Rotated_image and os.path.exists(Rotated_image):
            os.remove(Rotated_image)
            
        # 2. Only remove the original screenshot if it exists
        if os.path.exists(original_screenshot_path):
            os.remove(original_screenshot_path)
        
        try:
            input("Press Enter to exit....") 
        except EOFError:
            pass

if __name__=="__main__":
    file_path="newdocs1.docx"
    idle_timeOut_seconds=20 
    rotation_angle=0 #Quadrantal angles are suggested for image rotation ie,[0,90,180,270]
    # Note: +ve angles rotates image counter clock-wise and -ve angles rotate image clock-wise
    docx_width = 6.45
    docx_height = 9.67
    Screenshot_Automation(file_path,idle_timeOut_seconds,rotation_angle,docx_width,docx_height)